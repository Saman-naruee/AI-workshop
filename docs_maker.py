from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()

# Set title style
title = doc.add_heading("رزومه سامان ناروئی", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Contact Information
doc.add_paragraph("📧 samannaruee@gmail.com | ☎️ ۰۹۹۰۶۰۵۲۷۹۰")
doc.add_paragraph("📍 زاهدان، ایران | 🌐 github.com/Saman-naruee | 📎 linkedin.com/in/saman-naruee-nosrati")

# Job Title
doc.add_heading("برنامه‌نویس Python (توسعه‌دهنده API) | آماده برای همکاری دورکاری", level=2)

# Professional Summary
doc.add_heading("🧭 پروفایل حرفه‌ای", level=3)
doc.add_paragraph(
    "برنامه‌نویس بک‌اند با تمرکز بر توسعه سرویس‌های مقیاس‌پذیر و RESTful API با استفاده از Django. "
    "تجربه در طراحی سیستم‌های واقعی برای کسب‌و‌کارها از جمله همکاری با تیم دیجی‌کالا. مسلط بر Git، Docker، Redis و PostgreSQL. "
    "علاقه‌مند به یادگیری و استفاده از ابزارهای هوش مصنوعی مانند OpenRouter, Groq, Cohere, Copilot و Google Gemini در توسعه و ارتقاء کد."
)

# Skills
doc.add_heading("🧰 مهارت‌ها", level=3)
skills = [
    "زبان‌ها و فریم‌ورک‌ها: Python, Django, DRF",
    "ابزارها: Git, Docker, Celery, Redis, Postman",
    "دیتابیس‌ها: PostgreSQL, SQLite",
    "هوش مصنوعی و ابزارهای کدنویسی مدرن: OpenRouter, Groq, Github Copilot, Cohere, Google Gemini",
    "اصول توسعه: RESTful API Design, Clean Code, JWT Auth",
    "دیگر: کار تیمی، Agile, تجربه دورکاری"
]
for skill in skills:
    doc.add_paragraph(f"- {skill}", style='List Bullet')

# Work Experience
doc.add_heading("💼 تجربه‌های حرفه‌ای", level=3)

doc.add_paragraph("🔹 پروژه لندینگ دیجی‌کالا (کارآموزی کاریار استودیو - ۱۴۰۳)", style='List Bullet')
doc.add_paragraph(
    "• طراحی APIهای نسخه‌پذیر با DRF\n"
    "• پیاده‌سازی تاریخچه تغییرات با Django Simple History\n"
    "• اعتبارسنجی هوشمند داده‌ها و مجوزهای سفارشی\n"
    "🔗 github.com/karyar-studio/digimehr-landing-backend"
)

doc.add_paragraph("🔹 پروژه فروشگاه اینترنتی (خودآموز - ۱۴۰۲)", style='List Bullet')
doc.add_paragraph(
    "• طراحی کامل API محصولات، سفارشات، کاربران\n"
    "• مستندسازی با Swagger، احراز هویت JWT\n"
    "• بهینه‌سازی کوئری‌ها با ORM\n"
    "🔗 github.com/Saman-naruee/OnlineShopKaryar"
)

# Courses
doc.add_heading("🎓 دوره‌های تخصصی", level=3)
doc.add_paragraph("• گواهی Python مقدماتی – کوئرا")
doc.add_paragraph("• گواهی Python پیشرفته – کوئرا")
doc.add_paragraph("🔗 quera.org/user/SamanNaruee")

# Soft Skills
doc.add_heading("🌟 ویژگی‌های فردی", level=3)
soft_skills = [
    "مشتاق یادگیری مداوم",
    "روحیه تیمی بالا",
    "علاقه‌مند به هوش مصنوعی و فناوری‌های نوین"
]
for skill in soft_skills:
    doc.add_paragraph(f"- {skill}", style='List Bullet')

# Save the document
doc_path = "/mnt/data/Saman-Naruee-Customized-Resume.docx"
doc.save(doc_path)
doc_path
